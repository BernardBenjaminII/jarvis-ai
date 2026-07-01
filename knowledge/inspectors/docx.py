from pathlib import Path

from docx import Document

from .base import BaseInspector


class DocxInspector(BaseInspector):

    extensions = (".docx",)

    def inspect(self, path: Path):

        doc = Document(path)

        text = "\n".join(
            p.text
            for p in doc.paragraphs
        )

        return {

            "paragraphs": len(doc.paragraphs),

            "words": len(text.split()),

            "characters": len(text),
        }
